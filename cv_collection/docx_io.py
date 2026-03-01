"""
Helpers for reading CV content from Word documents.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document


def docx_to_text(path: Path) -> str | None:
    if path.name.startswith(".") or not zipfile.is_zipfile(path):
        return None
    try:
        document = Document(path)
    except Exception as e:
        print(f"⚠️  Cannot read {path}: {e}")
        return None

    para_map = {id(p._element): p for p in document.paragraphs}
    table_map = {id(t._element): t for t in document.tables}

    parts: list[str] = []
    for child in document.element.body:
        child_id = id(child)
        if child_id in para_map:
            txt = para_map[child_id].text.strip()
            if txt:
                parts.append(txt)
        elif child_id in table_map:
            for row in table_map[child_id].rows:
                seen: set[str] = set()
                cells: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen:
                        cells.append(cell_text)
                        seen.add(cell_text)
                if cells:
                    parts.append(" | ".join(cells))

    return "\n".join(parts) if parts else None
