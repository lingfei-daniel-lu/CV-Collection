"""
Shared helpers for parsing CV files and writing results.
"""

from __future__ import annotations

import json, zipfile
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document


def docx_to_text(path: Path) -> str | None:
    if path.name.startswith(".") or not zipfile.is_zipfile(path):
        return None
    try:
        document = Document(path)
    except Exception as e:
        print(f"⚠️  Cannot read {path}: {e}")
        return None
    para_map = {id(p._element): p for p in document.paragraphs}
    table_map = {id(t._element): t for t in document.tables}

    parts: list[str] = []
    for child in document.element.body:
        child_id = id(child)
        if child_id in para_map:
            txt = para_map[child_id].text.strip()
            if txt:
                parts.append(txt)
        elif child_id in table_map:
            for row in table_map[child_id].rows:
                seen: set[str] = set()
                cells: list[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen:
                        cells.append(cell_text)
                        seen.add(cell_text)
                if cells:
                    parts.append(" | ".join(cells))

    return "\n".join(parts) if parts else None


def safe_json_load(raw: str, *, label: str):
    if not raw:
        print(f"⚠️  Empty response for {label}")
        return None
    txt = raw.strip()
    if txt.startswith("```"):
        parts = txt.split("```", 2)
        if len(parts) >= 2:
            txt = parts[1]
            if txt.lower().startswith("json"):
                txt = txt[4:]
    start, end = txt.find("{"), txt.rfind("}")
    if start == -1 or end == -1:
        print(f"⚠️  No JSON braces in response for {label}")
        return None
    try:
        return json.loads(txt[start : end + 1])
    except json.JSONDecodeError as e:
        print(f"⚠️  JSON decode failed for {label}: {e}")
        return None


def flush_rows_to_csv(rows: list[dict], out_csv: Path, journals: Iterable[str]):
    if not rows:
        return
    df = pd.DataFrame(rows)
    first_cols = [
        "file",
        "name",
        "promotion_year",
        "promotion_university",
        "years_post_phd",
    ]
    journal_cols = list(journals)
    if not df.empty:
        df = df[first_cols + journal_cols]
    write_header = not out_csv.exists()
    df.to_csv(out_csv, index=False, mode="a", header=write_header)
